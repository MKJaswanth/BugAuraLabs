from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"C:\Users\jaswa\OneDrive\Documents\New project\output\C2C_Agri_Test_Case_Portfolio_Sample.docx"


COLORS = {
    "green": "2F6B3F",
    "light_green": "EAF4EC",
    "pale": "F7FAF7",
    "gray": "F2F4F5",
    "dark": "1F2933",
    "muted": "667085",
    "border": "D6E2D8",
    "amber": "FFF4D6",
}


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def borders(cell, color="D6E2D8", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def cell_margin(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_width(cell, width_inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, bold=False, color=None, size=8.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Aptos"
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(13 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.name = "Aptos Display"
    run.font.color.rgb = RGBColor.from_string(COLORS["green"] if level == 1 else COLORS["dark"])
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 12)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(COLORS["dark"])
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(9.5)
    return p


def add_callout(doc, title, body, fill="EAF4EC"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_width(cell, 6.6)
    shade(cell, fill)
    borders(cell, COLORS["border"])
    cell_margin(cell, 140, 170, 140, 170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(COLORS["green"])
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08
    r2 = p2.add_run(body)
    r2.font.name = "Aptos"
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor.from_string(COLORS["dark"])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_meta_table(doc):
    rows = [
        ("Project Type", "C2C agriculture marketplace platform"),
        ("Document Type", "Sanitized QA test case sample for portfolio use"),
        ("Testing Focus", "Functional, UI, validation, workflow, and basic negative testing"),
        ("Prepared By", "QA Tester / Manual Test Engineer"),
        ("Confidentiality", "All private names, links, contacts, IDs, credentials, pricing, and internal references removed"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (k, v) in enumerate(rows):
        left, right = table.rows[i].cells
        set_width(left, 1.85)
        set_width(right, 4.75)
        for c in (left, right):
            borders(c)
            cell_margin(c)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(left, COLORS["light_green"])
        set_cell_text(left, k, bold=True, color=COLORS["green"], size=9)
        set_cell_text(right, v, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_summary_table(doc):
    headers = ["Area", "Included"]
    rows = [
        ("Registration/Login", "Buyer and seller sign-up, login, logout, required field checks"),
        ("Product Listings", "Create, edit, view, search, filter, and validate crop/product posts"),
        ("Buyer Flow", "Browse listings, view details, contact seller, send enquiry/request"),
        ("Seller Flow", "Manage listing status, respond to buyer enquiry, update availability"),
        ("UI & Validation", "Error messages, mandatory fields, invalid formats, responsive checks"),
        ("Non-Functional Notes", "Smoke coverage for page load, session behavior, and usability observations"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_width(cell, 2.0 if j == 0 else 4.6)
        shade(cell, COLORS["green"])
        borders(cell, COLORS["green"])
        cell_margin(cell)
        set_cell_text(cell, h, bold=True, color="FFFFFF", size=9.5)
    for area, inc in rows:
        cells = table.add_row().cells
        for c in cells:
            borders(c)
            cell_margin(c)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cells[0], COLORS["pale"])
        set_cell_text(cells[0], area, bold=True, color=COLORS["green"], size=9)
        set_cell_text(cells[1], inc, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_test_cases(doc):
    headers = ["TC ID", "Module", "Scenario", "Steps", "Expected Result", "Priority", "Status"]
    widths = [0.58, 0.88, 1.42, 1.95, 1.72, 0.62, 0.58]
    cases = [
        ("TC-001", "Auth", "Register as buyer with valid details", "Open registration page; select Buyer; enter valid name, mobile/email, password, and location; submit.", "Buyer account is created and user lands on buyer dashboard or verification page.", "High", "Pass"),
        ("TC-002", "Auth", "Prevent registration with missing mandatory fields", "Open registration page; leave required fields blank; submit.", "Form displays clear required-field messages and does not create account.", "High", "Pass"),
        ("TC-003", "Auth", "Login with valid credentials", "Open login page; enter registered user credentials; submit.", "User is authenticated and redirected to the correct dashboard.", "High", "Pass"),
        ("TC-004", "Auth", "Reject login with invalid password", "Enter valid user ID/email with incorrect password; submit.", "System shows invalid credential message without revealing account details.", "High", "Pass"),
        ("TC-005", "Listings", "Seller creates new crop listing", "Login as seller; choose Add Product; enter crop name, quantity, price range, location, image, and description; submit.", "Listing is saved and visible in seller inventory with active/pending status.", "High", "Pass"),
        ("TC-006", "Listings", "Validate crop listing required fields", "Open Add Product; omit crop name, quantity, and location; submit.", "System highlights missing required fields and blocks submission.", "High", "Pass"),
        ("TC-007", "Listings", "Upload allowed product image", "Attach JPG/PNG image within allowed size while creating listing.", "Image uploads successfully and preview appears before/save after listing creation.", "Medium", "Pass"),
        ("TC-008", "Listings", "Reject unsupported image type", "Try uploading unsupported file type for listing image.", "System displays file type validation message and does not upload the file.", "Medium", "Pass"),
        ("TC-009", "Search", "Search listings by crop name", "Login as buyer; enter crop keyword in search field; apply search.", "Relevant crop listings are displayed and unrelated listings are reduced/hidden.", "High", "Pass"),
        ("TC-010", "Search", "Filter listings by location/category", "Apply category and location filters from listing page.", "Results refresh according to selected filters and selected filter state remains visible.", "Medium", "Pass"),
        ("TC-011", "Buyer Flow", "Open listing details page", "Select a product card from buyer listing page.", "Details page shows product name, quantity, seller area, description, and contact/enquiry action.", "High", "Pass"),
        ("TC-012", "Buyer Flow", "Send enquiry to seller", "Open listing details; click enquiry/contact action; enter message; submit.", "Enquiry is submitted and confirmation message is shown to buyer.", "High", "Pass"),
        ("TC-013", "Seller Flow", "Seller views buyer enquiries", "Login as seller; open enquiries section.", "Seller can view enquiry list with product reference, buyer message, and latest status.", "High", "Pass"),
        ("TC-014", "Seller Flow", "Update listing availability", "Open seller listing; change status from Available to Sold/Inactive; save.", "Updated status is reflected in seller dashboard and unavailable listing is handled correctly for buyers.", "High", "Pass"),
        ("TC-015", "Profile", "Update user profile details", "Open profile; edit address/location and contact preference; save.", "Profile changes are saved and visible after page refresh.", "Medium", "Pass"),
        ("TC-016", "Validation", "Validate numeric quantity and price fields", "Enter text/special characters in quantity or price fields; submit.", "System blocks invalid values and displays field-level validation.", "High", "Pass"),
        ("TC-017", "Session", "Logout from account", "Login; click logout; attempt to access dashboard using browser back button.", "User is logged out and protected pages require login again.", "High", "Pass"),
        ("TC-018", "Responsive", "Check listing page on mobile viewport", "Open buyer listing page on mobile screen size; review cards, filters, and actions.", "Content is readable, actions are accessible, and no important element overlaps.", "Medium", "Pass"),
        ("TC-019", "Usability", "Check empty search result state", "Search with a crop name that has no available listings.", "System shows a friendly no-results message and option to clear/change search.", "Low", "Pass"),
        ("TC-020", "Smoke", "Basic page load and navigation smoke test", "Open main pages: home, login, listing, details, profile/dashboard.", "Pages load without broken navigation, obvious console-visible UI failure, or missing critical content.", "High", "Pass"),
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_width(cell, widths[j])
        shade(cell, COLORS["green"])
        borders(cell, COLORS["green"])
        cell_margin(cell, 85, 70, 85, 70)
        set_cell_text(cell, h, bold=True, color="FFFFFF", size=7.8, align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, row in enumerate(cases):
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cell = cells[j]
            set_width(cell, widths[j])
            shade(cell, "FFFFFF" if idx % 2 == 0 else COLORS["pale"])
            borders(cell)
            cell_margin(cell, 80, 65, 80, 65)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            align = WD_ALIGN_PARAGRAPH.CENTER if j in (0, 5, 6) else WD_ALIGN_PARAGRAPH.LEFT
            color = COLORS["green"] if j in (0, 5, 6) else COLORS["dark"]
            set_cell_text(cell, val, bold=j in (0, 5, 6), color=color, size=7.0, align=align)


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Sanitized portfolio sample | C2C Agri QA Test Cases | No confidential client data included")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def set_portrait(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)


def set_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)


def build():
    doc = Document()
    section = doc.sections[0]
    set_portrait(section)
    add_footer(section)
    doc.core_properties.title = "C2C Agri Marketplace - Sanitized Test Case Document"
    doc.core_properties.subject = "Portfolio-ready QA test case sample"
    doc.core_properties.author = "QA Tester"
    doc.core_properties.comments = "Sanitized sample with confidential data removed."

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(36)
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("C2C Agri Marketplace")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(COLORS["green"])

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    r = subtitle.add_run("Sanitized Test Case Document for Portfolio Proof of Work")
    r.font.name = "Aptos"
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(COLORS["dark"])

    add_meta_table(doc)
    add_callout(
        doc,
        "Confidential Data Removed",
        "This document is a sanitized sample based on a C2C agriculture marketplace testing workflow. Real client identity, platform URLs, user credentials, personal contacts, exact business data, internal IDs, screenshots, and private notes have been replaced with generic placeholders.",
        COLORS["amber"],
    )

    add_heading(doc, "Project Overview", 1)
    add_body(
        doc,
        "The application under test is a consumer-to-consumer agriculture marketplace where sellers can list farm products and buyers can search, review, and send enquiries. The QA objective was to validate the key marketplace workflows from registration through listing discovery and enquiry handling.",
    )

    add_heading(doc, "Testing Scope", 1)
    add_summary_table(doc)

    add_heading(doc, "Test Approach", 1)
    for item in [
        "Reviewed the main buyer and seller journeys and identified high-risk functional areas.",
        "Prepared positive, negative, validation, and smoke test scenarios for core workflows.",
        "Executed tests using anonymized accounts and non-production sample data.",
        "Recorded expected results, priority, and pass/fail status for portfolio-level reporting.",
    ]:
        add_bullet(doc, item)

    doc.add_section(WD_SECTION.NEW_PAGE)
    set_landscape(doc.sections[-1])
    add_footer(doc.sections[-1])
    add_heading(doc, "Detailed Test Cases", 1)
    add_body(
        doc,
        "The table below contains representative test cases. The data has been intentionally generalized so the document can be shared publicly without exposing private project information.",
    )
    add_test_cases(doc)

    doc.add_section(WD_SECTION.NEW_PAGE)
    set_portrait(doc.sections[-1])
    add_footer(doc.sections[-1])
    add_heading(doc, "Sample Test Data", 1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["Data Type", "Sanitized Example", "Notes"]
    widths = [1.45, 2.45, 2.7]
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]
        set_width(c, widths[j])
        shade(c, COLORS["green"])
        borders(c, COLORS["green"])
        cell_margin(c)
        set_cell_text(c, h, bold=True, color="FFFFFF", size=9)
    rows = [
        ("Buyer Account", "buyer_sample_01 / masked email", "No real personal details used"),
        ("Seller Account", "seller_sample_01 / masked mobile", "Credentials removed"),
        ("Crop Listing", "Tomato, Onion, Wheat, Rice", "Generic crop names only"),
        ("Location", "Sample District / Sample State", "Exact address removed"),
        ("Price/Quantity", "Sample values", "Real rates and stock data removed"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            set_width(cells[j], widths[j])
            borders(cells[j])
            cell_margin(cells[j])
            shade(cells[j], "FFFFFF" if j else COLORS["pale"])
            set_cell_text(cells[j], val, bold=(j == 0), color=COLORS["green"] if j == 0 else COLORS["dark"], size=9)

    add_heading(doc, "Public Sharing Checklist", 1)
    for item in [
        "Client/company name removed or replaced with generic project name.",
        "Live URLs, API endpoints, database names, and admin paths removed.",
        "Real user information, phone numbers, emails, addresses, and credentials removed.",
        "Screenshots and attachments excluded unless separately approved for public use.",
        "Document positioned as a sanitized portfolio sample, not as a client-owned deliverable.",
    ]:
        add_bullet(doc, item)

    doc.save(OUT)


if __name__ == "__main__":
    build()
